"""Resume text extraction for PDF, DOCX, and plain text uploads."""
from __future__ import annotations

import io
from typing import BinaryIO


class UnsupportedFileError(ValueError):
    pass


def _read_pdf(stream: BinaryIO) -> str:
    from pypdf import PdfReader

    reader = PdfReader(stream)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _read_docx(stream: BinaryIO) -> str:
    from docx import Document

    document = Document(stream)
    lines = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(uploaded_file) -> str:
    """Extract text from a Streamlit UploadedFile-like object."""
    name = (getattr(uploaded_file, "name", "") or "").lower()

    # Prefer getvalue() so repeated calls on the same upload don't return empty
    # bytes after the buffer position has been advanced by a previous read.
    if hasattr(uploaded_file, "getvalue"):
        data = uploaded_file.getvalue()
    else:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        data = uploaded_file.read()

    if not data:
        raise UnsupportedFileError("The uploaded file is empty.")

    stream = io.BytesIO(data)

    if name.endswith(".pdf"):
        text = _read_pdf(stream)
    elif name.endswith(".docx"):
        text = _read_docx(stream)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileError(
            "Unsupported file type. Upload a PDF, DOCX, TXT, or MD file."
        )

    cleaned = "\n".join(line.rstrip() for line in text.splitlines() if line.strip())
    if not cleaned.strip():
        raise UnsupportedFileError(
            "Could not extract readable text. Try exporting the resume as a text-based PDF or DOCX."
        )
    return cleaned
